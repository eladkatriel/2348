"""
garmushka_word_injector.py
==========================
מקבל תוצאות של garmushka_pdf_analyzer ומכניס אותן לקובץ Word.

Placeholders שיש להוסיף לתבנית ה-Word:
  {{YearBuilt}}        ← שנת הגשה / בנייה
  {{TypicalFloor}}     ← תמונת תכנית קומה טיפוסית
  {{GroundFloor}}      ← תמונת תכנית קומת קרקע  (אופציונלי)
  {{BuildingSection}}  ← תמונת חתך

שימוש בקובץ app_with_garmushka.py:
--------------------------------------
from garmushka_pdf_analyzer import analyze_building_pdf
from garmushka_word_injector import inject_into_word

# בתוך process_item(), אחרי יצירת תיקיית הפרויקט:
pdf_data = find_and_analyze_garmushka_pdf(findings_folder_local_or_dropbox)
if pdf_data:
    inject_into_word(doc, pdf_data)   # doc = python-docx Document object
"""

import io
import os
import tempfile
from typing import Optional
from docx import Document
from docx.shared import Inches


# גדלי תמונות ברירת מחדל בתוך Word (אינצ'ים)
DEFAULT_PLAN_WIDTH   = 6.0   # תכנית קומה
DEFAULT_SECTION_WIDTH = 6.0  # חתך


def inject_into_word(doc: Document, pdf_data: dict) -> bool:
    """
    מכניס נתוני גרמושקא לתוך Document פתוח.
    מחזיר True אם הוכנס לפחות שדה אחד.

    pdf_data = תוצאה של analyze_building_pdf()
    """
    injected = False

    # ── תאריך / שנה ──────────────────────────────────────────────────────────
    date_val = pdf_data.get("date") or ""
    year_val = date_val[:4] if len(date_val) >= 4 else date_val
    if year_val:
        _replace_text_placeholder(doc, "{{YearBuilt}}", year_val)
        _replace_text_placeholder(doc, "{{DateBuilt}}", date_val)
        injected = True
        print(f"[Injector] תאריך: {date_val}")

    # ── תכנית קומה טיפוסית ───────────────────────────────────────────────────
    if pdf_data.get("typical_floor_image"):
        ok = _replace_image_placeholder(
            doc, "{{TypicalFloor}}",
            pdf_data["typical_floor_image"],
            width_inches=DEFAULT_PLAN_WIDTH,
        )
        if ok:
            injected = True
            print("[Injector] תכנית קומה טיפוסית הוכנסה")
        else:
            print("[Injector] לא נמצא {{TypicalFloor}} בתבנית")

    # ── תכנית קומת קרקע ──────────────────────────────────────────────────────
    if pdf_data.get("ground_floor_image"):
        ok = _replace_image_placeholder(
            doc, "{{GroundFloor}}",
            pdf_data["ground_floor_image"],
            width_inches=DEFAULT_PLAN_WIDTH,
        )
        if ok:
            injected = True
            print("[Injector] תכנית קומת קרקע הוכנסה")

    # ── חתך ──────────────────────────────────────────────────────────────────
    if pdf_data.get("section_image"):
        ok = _replace_image_placeholder(
            doc, "{{BuildingSection}}",
            pdf_data["section_image"],
            width_inches=DEFAULT_SECTION_WIDTH,
        )
        if ok:
            injected = True
            print("[Injector] חתך הוכנס")
        else:
            print("[Injector] לא נמצא {{BuildingSection}} בתבנית")

    return injected


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _replace_text_placeholder(doc: Document, placeholder: str, value: str):
    """מחליף placeholder טקסטואלי בכל חלקי המסמך."""
    def _in_para(para):
        full = "".join(r.text for r in para.runs)
        if placeholder in full:
            new = full.replace(placeholder, value)
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        _in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _in_para(para)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for para in part.paragraphs:
                _in_para(para)


def _replace_image_placeholder(
    doc: Document,
    placeholder: str,
    image_path: str,
    width_inches: float = 6.0,
) -> bool:
    """
    מחליף placeholder בתמונה.
    מחזיר True אם נמצא ה-placeholder והוחלף.
    """
    if not os.path.exists(image_path):
        print(f"[Injector] קובץ תמונה לא נמצא: {image_path}")
        return False

    def _try_insert(para) -> bool:
        full = "".join(r.text for r in para.runs)
        if placeholder not in full:
            return False
        # נקה את ה-placeholder
        for r in para.runs:
            r.text = r.text.replace(placeholder, "")
        # הוסף תמונה ב-run חדש
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        return True

    for para in doc.paragraphs:
        if _try_insert(para):
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _try_insert(para):
                        return True
    for section in doc.sections:
        for part in (section.header, section.footer):
            for para in part.paragraphs:
                if _try_insert(para):
                    return True
    return False


# ─────────────────────────────────────────────────────────────────────────────
# Dropbox integration: find garmushka PDFs in project folder
# ─────────────────────────────────────────────────────────────────────────────

def find_garmushka_pdfs_in_dropbox(dbx, folder_path: str) -> list:
    """
    מחפש קבצי PDF בתיקיית הפרויקט ב-Dropbox שנראים כתיקי בניין.
    מחזיר רשימה של נתיבי Dropbox.
    """
    TRIGGER_WORDS = ["היתר", "הגשה", "גרמושקא", "גרמושקה", "תיק", "permit", "building"]
    found = []
    try:
        result = dbx.files_list_folder(folder_path, recursive=True)
        while True:
            for entry in result.entries:
                name_lower = entry.name.lower()
                if name_lower.endswith(".pdf"):
                    if any(w in name_lower for w in [w.lower() for w in TRIGGER_WORDS]):
                        found.append(entry.path_display)
                        print(f"[Finder] נמצא PDF: {entry.name}")
            if result.has_more:
                result = dbx.files_list_folder_continue(result.cursor)
            else:
                break
    except Exception as e:
        print(f"[Finder] שגיאת Dropbox: {e}")
    return found


def download_and_analyze(dbx, dropbox_pdf_path: str, tmp_dir: str) -> Optional[dict]:
    """
    מוריד PDF מ-Dropbox, מנתח אותו, מחזיר dict.
    """
    from garmushka_pdf_analyzer import analyze_building_pdf
    import os

    local_path = os.path.join(tmp_dir, os.path.basename(dropbox_pdf_path))
    try:
        _, res = dbx.files_download(dropbox_pdf_path)
        with open(local_path, "wb") as f:
            f.write(res.content)
        print(f"[Downloader] הורד: {os.path.basename(dropbox_pdf_path)}")
        return analyze_building_pdf(local_path, output_dir=tmp_dir)
    except Exception as e:
        print(f"[Downloader] שגיאה: {e}")
        return None
