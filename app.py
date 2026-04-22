import io
import os
import re
import json
import math
from datetime import datetime
from pathlib import PurePosixPath

import requests
import dropbox
from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Inches

# ── Garmushka PDF ─────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable, Image as RLImage, Paragraph,
    SimpleDocTemplate, Spacer, Table, TableStyle,
)

_G_BLUE_DARK  = colors.HexColor("#0C447C")
_G_BLUE_MID   = colors.HexColor("#378ADD")
_G_BLUE_LIGHT = colors.HexColor("#E6F1FB")
_G_GRAY_BG    = colors.HexColor("#F5F5F3")
_G_GRAY_BDR   = colors.HexColor("#D3D1C7")
_G_GRAY_TXT   = colors.HexColor("#444441")
_G_GRAY_MUT   = colors.HexColor("#888780")
_G_WHITE      = colors.white
_G_PAGE_W, _G_PAGE_H = A4
_G_MARGIN     = 2 * cm
_G_FONT_NAME  = "Helvetica"   # ייעודי לגרמושקא — מוחלף ב-_g_init_font()

def _g_init_font() -> str:
    global _G_FONT_NAME
    for path in [
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    ]:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("GHEB", path))
                _G_FONT_NAME = "GHEB"
                return _G_FONT_NAME
            except Exception:
                pass
    return _G_FONT_NAME

_g_init_font()

def _g_bidi(text: str) -> str:
    """RTL rendering for Hebrew text in reportlab."""
    if not text:
        return ""
    try:
        from bidi.algorithm import get_display
        return get_display(str(text))
    except ImportError:
        s = str(text)
        return s[::-1] if any("\u0590" <= c <= "\u05FF" for c in s) else s

def _g_styles() -> dict:
    f = _G_FONT_NAME
    return {
        "title":    ParagraphStyle("gt",  fontName=f, fontSize=20, textColor=_G_BLUE_DARK,  alignment=2, leading=28),
        "subtitle": ParagraphStyle("gs",  fontName=f, fontSize=13, textColor=_G_BLUE_DARK,  alignment=2, leading=20),
        "section":  ParagraphStyle("gsc", fontName=f, fontSize=12, textColor=_G_BLUE_DARK,  alignment=2, leading=18, spaceBefore=10, spaceAfter=6),
        "label":    ParagraphStyle("gl",  fontName=f, fontSize=9,  textColor=_G_GRAY_MUT,   alignment=2, leading=13),
        "value":    ParagraphStyle("gv",  fontName=f, fontSize=13, textColor=_G_GRAY_TXT,   alignment=2, leading=18),
        "small":    ParagraphStyle("gsm", fontName=f, fontSize=8,  textColor=_G_GRAY_MUT,   alignment=2, leading=12),
        "mfb":      ParagraphStyle("gmf", fontName=f, fontSize=10, textColor=_G_GRAY_MUT,   alignment=1, leading=16),
    }

def _g_page_frame(c, doc):
    c.saveState()
    c.setStrokeColor(colors.HexColor("#B5D4F4"))
    c.setLineWidth(0.5)
    c.rect(_G_MARGIN*0.6, _G_MARGIN*0.6, _G_PAGE_W-_G_MARGIN*1.2, _G_PAGE_H-_G_MARGIN*1.2)
    c.setFont("Helvetica", 8)
    c.setFillColor(_G_GRAY_MUT)
    c.drawCentredString(_G_PAGE_W/2, _G_MARGIN*0.3, str(doc.page))
    c.restoreState()

def _g_meta_grid(story, st, cells: list):
    """מציג רשת של תאי label+value, שורות של 2."""
    col_w = (_G_PAGE_W - 2*_G_MARGIN) / 2 - 4
    for i in range(0, len(cells), 2):
        pair = cells[i:i+2]
        while len(pair) < 2:
            pair.append(("", ""))
        cell_tables = []
        for label, value in pair:
            ct = Table(
                [[Paragraph(_g_bidi(label), st["label"])],
                 [Paragraph(_g_bidi(str(value or "—")), st["value"])]],
                colWidths=[col_w],
            )
            ct.setStyle(TableStyle([
                ("BACKGROUND",    (0,0),(-1,-1), _G_GRAY_BG),
                ("ALIGN",         (0,0),(-1,-1), "RIGHT"),
                ("TOPPADDING",    (0,0),(-1,-1), 10),
                ("BOTTOMPADDING", (0,0),(-1,-1), 10),
                ("LEFTPADDING",   (0,0),(-1,-1), 12),
                ("RIGHTPADDING",  (0,0),(-1,-1), 12),
            ]))
            cell_tables.append(ct)
        row_tbl = Table([cell_tables], colWidths=[col_w+6, col_w+6])
        row_tbl.setStyle(TableStyle([
            ("LEFTPADDING",  (0,0),(-1,-1), 2),
            ("RIGHTPADDING", (0,0),(-1,-1), 2),
            ("VALIGN",       (0,0),(-1,-1), "TOP"),
        ]))
        story.append(row_tbl)
        story.append(Spacer(1, 5))

def build_garmushka_pdf(data: dict, map_bytes: bytes | None = None) -> bytes:
    """
    יוצר גרמושקא PDF ב-memory ומחזיר bytes.
    data = אותו dict שמגיע מ-get_item_data() (עמודות Monday).
    map_bytes = bytes של תמונת המפה שכבר נוצרה, או None.
    """
    buf = io.BytesIO()
    st  = _g_styles()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=_G_MARGIN, leftMargin=_G_MARGIN,
        topMargin=_G_MARGIN,   bottomMargin=_G_MARGIN,
        title="גרמושקא — נסח מבנה",
    )

    # ── שליפת שדות ───────────────────────────────────────────────────────────
    street    = (data.get("text_mm12vcy9",  "") or "").strip()
    number    = (data.get("text_mm12jf0w",  "") or "").strip()
    apartment = (data.get("text_mm127a33",  "") or "").strip()
    city      = (data.get("text_mm264acy",  "") or "").strip()
    engineer  = (data.get("text_mm12tcvb",  "") or "").strip()
    case      = (data.get("text_mm12qp1q",  "") or "").strip()
    id_num    = (data.get("text_mm12vayb",  "") or "").strip()
    months    = (data.get("text_mm129ef8",  "") or "").strip()
    hit_date  = format_yyyy_mm_dd_to_dd_mm_yyyy((data.get("date_mm1rnmvg", "") or "").strip())
    rep_type  = (data.get(REPORT_TYPE_COLUMN_ID, "") or "").strip()

    apt_part = f"דירה {apartment}" if apartment else ""
    address  = " ".join(p for p in [f"רחוב {street}", number, apt_part, city] if p)

    story = []

    # ── כותרת ────────────────────────────────────────────────────────────────
    header = Table(
        [[Paragraph(_g_bidi("גרמושקא — נסח מבנה"), st["title"])],
         [Paragraph(_g_bidi(address), st["subtitle"])]],
        colWidths=[_G_PAGE_W - 2*_G_MARGIN],
    )
    header.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,-1), _G_BLUE_LIGHT),
        ("ALIGN",         (0,0),(-1,-1), "RIGHT"),
        ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
        ("TOPPADDING",    (0,0),(-1,-1), 14),
        ("BOTTOMPADDING", (0,0),(-1,-1), 14),
        ("LEFTPADDING",   (0,0),(-1,-1), 18),
        ("RIGHTPADDING",  (0,0),(-1,-1), 18),
    ]))
    story.append(header)
    story.append(Spacer(1, 0.4*cm))

    # ── נתוני מבנה ───────────────────────────────────────────────────────────
    story.append(Paragraph(_g_bidi("נתוני מבנה"), st["section"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=_G_BLUE_MID, spaceAfter=8))
    _g_meta_grid(story, st, [
        ("רחוב",         f"{street} {number}".strip()),
        ("דירה",          apartment or "—"),
        ("עיר",           city or "—"),
        ("תאריך פגיעה",   hit_date or "—"),
    ])
    story.append(Spacer(1, 0.3*cm))

    # ── נתוני תיק ────────────────────────────────────────────────────────────
    story.append(Paragraph(_g_bidi("נתוני תיק"), st["section"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=_G_BLUE_MID, spaceAfter=8))
    _g_meta_grid(story, st, [
        ("מהנדס",       engineer or "—"),
        ("תיק",         case     or "—"),
        ("ת.ז / ח.פ",  id_num   or "—"),
        ("חודשים",      months   or "—"),
    ])
    if rep_type:
        story.append(Spacer(1, 4))
        rt = Table(
            [[Paragraph(_g_bidi("סוג דו\"ח"), st["label"]),
              Paragraph(_g_bidi(rep_type),    st["value"])]],
            colWidths=[(_G_PAGE_W-2*_G_MARGIN)*0.3, (_G_PAGE_W-2*_G_MARGIN)*0.7],
        )
        rt.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(-1,-1), _G_GRAY_BG),
            ("ALIGN",         (0,0),(-1,-1), "RIGHT"),
            ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
            ("TOPPADDING",    (0,0),(-1,-1), 8),
            ("BOTTOMPADDING", (0,0),(-1,-1), 8),
            ("LEFTPADDING",   (0,0),(-1,-1), 10),
            ("RIGHTPADDING",  (0,0),(-1,-1), 10),
        ]))
        story.append(rt)
    story.append(Spacer(1, 0.4*cm))

    # ── מפה ──────────────────────────────────────────────────────────────────
    story.append(Paragraph(_g_bidi("מפה"), st["section"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=_G_BLUE_MID, spaceAfter=8))
    avail_w  = _G_PAGE_W - 2*_G_MARGIN
    map_h    = 9*cm
    if map_bytes:
        try:
            img = RLImage(io.BytesIO(map_bytes), width=avail_w, height=map_h)
            img.hAlign = "CENTER"
            story.append(img)
        except Exception as e:
            print("GARMUSHKA MAP IMAGE FAILED:", e)
            map_bytes = None   # fall through to text fallback
    if not map_bytes:
        fb = Table(
            [[Paragraph(_g_bidi(MAP_NOT_FOUND_TEXT), st["mfb"])]],
            colWidths=[avail_w], rowHeights=[map_h],
        )
        fb.setStyle(TableStyle([
            ("BACKGROUND", (0,0),(-1,-1), _G_GRAY_BG),
            ("ALIGN",      (0,0),(-1,-1), "CENTER"),
            ("VALIGN",     (0,0),(-1,-1), "MIDDLE"),
            ("BOX",        (0,0),(-1,-1), 0.5, _G_GRAY_BDR),
        ]))
        story.append(fb)

    # ── כותרת תחתונה ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.6*cm))
    story.append(HRFlowable(width="100%", thickness=0.3, color=_G_GRAY_MUT, spaceAfter=4))
    story.append(Paragraph(_g_bidi(f"תאריך הפקה: {datetime.now().strftime('%d.%m.%Y')}"), st["small"]))
    story.append(Paragraph(
        _g_bidi("המידע נשאב ממאגרי מידע ממשלתיים ועירוניים. יש לאמת מול הרשות המקומית לפני שימוש משפטי."),
        st["small"],
    ))

    doc.build(story, onFirstPage=_g_page_frame, onLaterPages=_g_page_frame)
    return buf.getvalue()

app = Flask(__name__)

MONDAY_API_KEY = os.environ.get("MONDAY_API_KEY")
BOARD_ID = os.environ.get("BOARD_ID")
DROPBOX_APP_KEY = os.environ.get("DROPBOX_APP_KEY")
DROPBOX_APP_SECRET = os.environ.get("DROPBOX_APP_SECRET")
DROPBOX_REFRESH_TOKEN = os.environ.get("DROPBOX_REFRESH_TOKEN")
DROPBOX_TOKEN = os.environ.get("DROPBOX_TOKEN")
MAPBOX_ACCESS_TOKEN = os.environ.get("MAPBOX_ACCESS_TOKEN")

LINK_COLUMN_ID = os.environ.get("LINK_COLUMN_ID", "link_mm27m1ce").strip()
FILES_COLUMN_ID = "FILES"
DROPBOX_SHARED_LINK = os.environ.get("DROPBOX_SHARED_LINK", "https://www.dropbox.com/scl/fo/5vychlhm7el3kjn5t8ah9/h?rlkey=fzledyaec01ixjsnd1zgbqoax&st=o32icz46&dl=0").strip()

TARGET_REPORTS_FOLDER_NAME = "20260228 - שאגת הארי"
TEMPLATE_RELATIVE_DIR = "Template/23-48"
CITY_COLUMN_ID = "text_mm264acy"
CASE_COLUMN_ID = "text_mm12qp1q"
ID_COLUMN_ID = "text_mm12vayb"
REPORT_TYPE_COLUMN_ID = "color_mm12nfzt"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

COLUMN_MAP = {
    "{{Engineer}}": "text_mm12tcvb",
    "{{Street}}": "text_mm12vcy9",
    "{{Number}}": "text_mm12jf0w",
    "{{Apartment}}": "text_mm127a33",
    "{{Hit date}}": "date_mm1rnmvg",
    "{{months}}": "text_mm129ef8",
    "{{City}}": CITY_COLUMN_ID,
    "{{Case}}": CASE_COLUMN_ID,
    "{{ID}}": ID_COLUMN_ID,
    "{{Name}}": "text_mm12yqvr",
}

MAP_WIDTH_INCHES = 6.5
MAP_HEIGHT_INCHES = 4.4
MAP_IMAGE_WIDTH = 640
MAP_IMAGE_HEIGHT = 434
TARGET_HORIZONTAL_METERS = 1000.0
MAPBOX_GEOCODE_URL = "https://api.mapbox.com/search/geocode/v6/forward"
MAPBOX_STYLE_PATH = "mapbox/satellite-streets-v12"
MAPBOX_STATIC_BASE = f"https://api.mapbox.com/styles/v1/{MAPBOX_STYLE_PATH}/static"
MAP_MARKER_MODE = os.environ.get("MAP_MARKER_MODE", "small").strip().lower()
MAP_NOT_FOUND_TEXT = "לא נמצא מפה תואמת לכתובת"
MAP_CACHE = {}

if not MONDAY_API_KEY:
    raise ValueError("Missing MONDAY_API_KEY environment variable")
if not BOARD_ID:
    raise ValueError("Missing BOARD_ID environment variable")

def init_dropbox():
    if DROPBOX_REFRESH_TOKEN:
        if not DROPBOX_APP_KEY or not DROPBOX_APP_SECRET:
            raise ValueError("Using DROPBOX_REFRESH_TOKEN requires DROPBOX_APP_KEY and DROPBOX_APP_SECRET")
        return dropbox.Dropbox(
            oauth2_refresh_token=DROPBOX_REFRESH_TOKEN,
            app_key=DROPBOX_APP_KEY,
            app_secret=DROPBOX_APP_SECRET,
            timeout=120,
        )
    if DROPBOX_TOKEN:
        return dropbox.Dropbox(DROPBOX_TOKEN, timeout=120)
    raise ValueError("Missing Dropbox credentials. Set DROPBOX_REFRESH_TOKEN + DROPBOX_APP_KEY + DROPBOX_APP_SECRET, or DROPBOX_TOKEN")

dbx = init_dropbox()
print("Dropbox initialized")

def path_exists(path: str) -> bool:
    try:
        dbx.files_get_metadata(path)
        return True
    except Exception as e:
        print("PATH NOT FOUND:", path, str(e))
        return False

def resolve_shared_root_from_link(shared_link: str) -> str:
    if not shared_link:
        raise Exception("DROPBOX_SHARED_LINK is missing")
    meta = dbx.sharing_get_shared_link_metadata(shared_link)
    shared_folder_id = getattr(meta, "shared_folder_id", None)
    name = getattr(meta, "name", "")
    print("SHARED LINK NAME:", name)
    print("SHARED LINK SHARED_FOLDER_ID:", shared_folder_id)

    if shared_folder_id:
        listing = dbx.sharing_list_folders(limit=100)
        while True:
            for folder in listing.entries:
                if str(getattr(folder, "shared_folder_id", "")) == str(shared_folder_id):
                    mounted_path = getattr(folder, "path_display", None) or getattr(folder, "path_lower", None)
                    if mounted_path:
                        print("RESOLVED SHARED ROOT FROM LINK:", mounted_path)
                        return mounted_path
            if getattr(listing, "has_more", False):
                listing = dbx.sharing_list_folders_continue(listing.cursor)
            else:
                break

    for attr_name in ("path_display", "path_lower"):
        value = getattr(meta, attr_name, None)
        if value:
            print("RESOLVED SHARED ROOT FROM LINK METADATA:", value)
            return value

    raise Exception("Could not resolve mounted Dropbox path from the shared link")

def resolve_base_reports_and_template_dir(shared_root: str):
    shared_root = shared_root.rstrip("/")
    base_name = PurePosixPath(shared_root).name
    child_reports = f"{shared_root}/{TARGET_REPORTS_FOLDER_NAME}"
    child_template_dir = f"{shared_root}/{TEMPLATE_RELATIVE_DIR}"
    direct_reports = shared_root
    template_candidates = [child_template_dir, "/Template/23-48"]

    if base_name == TARGET_REPORTS_FOLDER_NAME and path_exists(direct_reports):
        base_reports_path = direct_reports
    elif path_exists(child_reports):
        base_reports_path = child_reports
    else:
        raise Exception(f"Could not resolve the reports root from the shared link. Checked: {direct_reports} | {child_reports}")

    template_dir = None
    for candidate in template_candidates:
        if path_exists(candidate):
            template_dir = candidate
            break

    if not template_dir:
        raise Exception("Could not resolve template directory from the shared link context. Checked: " + " | ".join(template_candidates))

    print("FINAL BASE_REPORTS_PATH:", base_reports_path)
    print("FINAL TEMPLATE_DIR:", template_dir)
    return base_reports_path, template_dir

SHARED_ROOT_PATH = resolve_shared_root_from_link(DROPBOX_SHARED_LINK)
BASE_REPORTS_PATH, TEMPLATE_DIR = resolve_base_reports_and_template_dir(SHARED_ROOT_PATH)

def monday_query(query: str, variables=None):
    response = requests.post(
        "https://api.monday.com/v2",
        json={"query": query, "variables": variables or {}},
        headers={"Authorization": MONDAY_API_KEY, "Content-Type": "application/json", "API-Version": "2026-04"},
        timeout=60,
    )
    response.raise_for_status()
    payload = response.json()
    if "errors" in payload:
        raise Exception(f"monday API error: {payload['errors']}")
    return payload

def get_item_data(item_id: int):
    query = """
    query ($item_ids: [ID!]) {
      items(ids: $item_ids) {
        id
        name
        column_values {
          id
          text
        }
      }
    }
    """
    data = monday_query(query, {"item_ids": [str(item_id)]})
    items = data.get("data", {}).get("items", [])
    if not items:
        raise Exception(f"No item found for item_id={item_id}")
    item = items[0]
    cols = {c["id"]: c.get("text", "") for c in item["column_values"]}
    cols["name"] = item["name"]
    cols["item_id"] = str(item["id"])
    return cols

def update_link_column(item_id: int, column_id: str, url: str, text: str):
    column_values = json.dumps({column_id: {"url": url, "text": text}})
    query = """
    mutation ($board_id: ID!, $item_id: ID!, $column_values: JSON!) {
      change_multiple_column_values(board_id: $board_id, item_id: $item_id, column_values: $column_values) { id }
    }
    """
    result = monday_query(query, {"board_id": str(BOARD_ID), "item_id": str(item_id), "column_values": column_values})
    print("LINK COLUMN UPDATE RESULT:", result)
    return result

def upload_file_to_monday(item_id: int, column_id: str, file_name: str, file_bytes: bytes):
    query = """
    mutation ($item_id: ID!, $column_id: String!, $file: File!) {
      add_file_to_column(item_id: $item_id, column_id: $column_id, file: $file) { id }
    }
    """
    data = {"query": query, "variables": json.dumps({"item_id": str(item_id), "column_id": column_id})}
    files = {"variables[file]": (file_name, file_bytes, DOCX_MIME)}
    response = requests.post("https://api.monday.com/v2/file", headers={"Authorization": MONDAY_API_KEY, "API-Version": "2026-04"}, data=data, files=files, timeout=120)
    response.raise_for_status()
    payload = response.json()
    if "errors" in payload:
        raise Exception(f"monday file upload error: {payload['errors']}")
    return payload

def sanitize_filename(name: str) -> str:
    invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
    for ch in invalid_chars:
        name = name.replace(ch, "-")
    return " ".join(name.split()).strip()

def create_folder_if_needed(folder_path: str):
    try:
        dbx.files_create_folder_v2(folder_path)
        print("FOLDER CREATED:", folder_path)
    except Exception as e:
        print("FOLDER CREATE SKIPPED OR FAILED:", folder_path, str(e))

def get_or_create_shared_link(file_path: str) -> str:
    try:
        existing_links = dbx.sharing_list_shared_links(path=file_path, direct_only=True).links
        if existing_links:
            print("EXISTING SHARE LINK FOUND")
            return existing_links[0].url
    except Exception as e:
        print("SHARED LINK LOOKUP FAILED:", str(e))
    link = dbx.sharing_create_shared_link_with_settings(file_path).url
    print("NEW SHARE LINK CREATED:", link)
    return link

def format_yyyy_mm_dd_to_dd_mm_yyyy(date_str: str) -> str:
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return dt.strftime("%d.%m.%Y")
    except Exception:
        return date_str

def extract_first_date(value: str) -> str:
    if not value:
        return ""
    match = re.search(r"\d{4}-\d{2}-\d{2}", value)
    if not match:
        return value.strip()
    return format_yyyy_mm_dd_to_dd_mm_yyyy(match.group(0))

def resolve_template_filename(report_type_value: str) -> str:
    value = (report_type_value or "").strip()
    mapping = {
        "קבלן": "Contractor_template.docx",
        "מהנדס": "Engineer_template.docx",
        "להריסה": "Engineer_template.docx",
        "נזק ישן, אין פינוי ואין פיצוי": "Contractor_template.docx",
    }
    selected = mapping.get(value, "Contractor_template.docx")
    print("REPORT TYPE VALUE:", value)
    print("SELECTED TEMPLATE FILE:", selected)
    return selected

def build_template_path(report_type_value: str) -> str:
    template_filename = resolve_template_filename(report_type_value)
    template_path = f"{TEMPLATE_DIR}/{template_filename}"
    if not path_exists(template_path):
        raise Exception(f"Template not found at exact path: {template_path}")
    print("TEMPLATE PATH:", template_path)
    return template_path

def normalize_hebrew_for_match(text: str) -> str:
    text = (text or "").strip().lower()
    text = text.replace(",", " ")
    text = re.sub(r"\s+", " ", text)
    return text

def build_geocode_queries(street: str, number: str, city: str):
    street = (street or "").strip()
    number = (number or "").strip()
    city = (city or "").strip()
    queries = [
        f"{street} {number}, {city}, Israel",
        f"{street}, {number}, {city}, Israel",
        f"{street} {number}, {city}",
        f"{street}, {number}, {city}",
    ]
    if street.startswith("נחל "):
        alt_street = street.replace("נחל ", "", 1).strip()
        queries.extend([
            f"{alt_street} {number}, {city}, Israel",
            f"{alt_street}, {number}, {city}, Israel",
        ])
    deduped = []
    seen = set()
    for q in queries:
        key = normalize_hebrew_for_match(q)
        if key not in seen:
            seen.add(key)
            deduped.append(q)
    return deduped

def mapbox_forward_geocode(query_text: str):
    if not MAPBOX_ACCESS_TOKEN:
        raise Exception("Missing MAPBOX_ACCESS_TOKEN environment variable")
    params = {
        "q": query_text,
        "country": "il",
        "types": "address",
        "limit": 3,
        "language": "he",
        "access_token": MAPBOX_ACCESS_TOKEN,
    }
    response = requests.get(MAPBOX_GEOCODE_URL, params=params, timeout=30)
    response.raise_for_status()
    payload = response.json()
    return payload.get("features", [])

def validate_feature(feature: dict, expected_city: str, expected_street: str, expected_number: str) -> bool:
    expected_city_n = normalize_hebrew_for_match(expected_city)
    expected_street_n = normalize_hebrew_for_match(expected_street).replace("נחל ", "").strip()
    expected_number_n = normalize_hebrew_for_match(expected_number)

    props = feature.get("properties", {}) or {}
    full_address = normalize_hebrew_for_match(props.get("full_address", ""))
    name = normalize_hebrew_for_match(props.get("name", ""))
    context = props.get("context", {}) or {}

    place_name = normalize_hebrew_for_match((context.get("place") or {}).get("name", ""))
    street_name = normalize_hebrew_for_match((context.get("street") or {}).get("name", ""))
    address_number = normalize_hebrew_for_match((context.get("address") or {}).get("address_number", ""))

    street_match = (
        expected_street_n in full_address
        or expected_street_n in name
        or expected_street_n == street_name.replace("נחל ", "").strip()
    )
    city_match = (expected_city_n in full_address) or (expected_city_n == place_name)
    number_match = (expected_number_n in full_address) or (expected_number_n in name) or (expected_number_n == address_number)

    return street_match and city_match and number_match

def geocode_address_mapbox(street: str, number: str, city: str):
    queries = build_geocode_queries(street, number, city)
    print("MAPBOX GEOCODE QUERIES:", queries)
    all_candidates = []
    for query_text in queries:
        features = mapbox_forward_geocode(query_text)
        if not features:
            continue
        for feature in features:
            print("MAPBOX GEOCODE RESULT:", feature)
            all_candidates.append(feature)
            if validate_feature(feature, city, street, number):
                lon, lat = feature["geometry"]["coordinates"]
                print("MAPBOX GEOCODE VALIDATED LON/LAT:", lon, lat)
                return float(lon), float(lat)
    raise Exception(
        "Mapbox geocoding did not return a validated address match. "
        f"Expected street='{street}', number='{number}', city='{city}'. "
        f"Candidates checked: {len(all_candidates)}"
    )

def calc_zoom_for_target_width(lat: float, image_width_px: int, target_width_m: float) -> float:
    total_pixels = image_width_px * 2
    meters_per_pixel = target_width_m / total_pixels
    zoom = math.log2((156543.03392 * math.cos(math.radians(lat))) / meters_per_pixel)
    zoom = max(0.0, min(22.0, round(zoom, 2)))
    print("MAPBOX TARGET WIDTH METERS:", target_width_m)
    print("MAPBOX CALCULATED ZOOM:", zoom)
    return zoom

def build_mapbox_static_url(lon: float, lat: float, zoom: float) -> str:
    overlay = ""
    if MAP_MARKER_MODE == "small":
        overlay = f"pin-s+f00({lon},{lat})/"
    elif MAP_MARKER_MODE == "none":
        overlay = ""
    return (
        f"{MAPBOX_STATIC_BASE}/{overlay}{lon},{lat},{zoom}/"
        f"{MAP_IMAGE_WIDTH}x{MAP_IMAGE_HEIGHT}@2x"
        f"?access_token={MAPBOX_ACCESS_TOKEN}&logo=false&attribution=false"
    )

def create_mapbox_hybrid_image(street: str, number: str, city: str) -> bytes:
    cache_key = f"{street}|{number}|{city}|{MAP_MARKER_MODE}|{TARGET_HORIZONTAL_METERS}"
    if cache_key in MAP_CACHE:
        print("MAP CACHE HIT")
        return MAP_CACHE[cache_key]
    lon, lat = geocode_address_mapbox(street, number, city)
    zoom = calc_zoom_for_target_width(lat, MAP_IMAGE_WIDTH, TARGET_HORIZONTAL_METERS)
    url = build_mapbox_static_url(lon, lat, zoom)
    print("MAPBOX STYLE:", MAPBOX_STYLE_PATH)
    print("MAPBOX MARKER MODE:", MAP_MARKER_MODE)
    print("MAPBOX STATIC URL:", url.replace(MAPBOX_ACCESS_TOKEN, "***"))
    response = requests.get(url, timeout=60)
    response.raise_for_status()
    image_bytes = response.content
    MAP_CACHE[cache_key] = image_bytes
    return image_bytes

def replace_in_paragraph(paragraph, replacements: dict):
    if not paragraph.text:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return
    new_text = full_text
    changed = False
    for old, new in replacements.items():
        if old in new_text:
            new_text = new_text.replace(old, str(new))
            changed = True
    if not changed:
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""

def replace_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            replace_in_document_parts(cell, replacements)

def replace_in_document_parts(container, replacements: dict):
    for paragraph in container.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in container.tables:
        replace_in_table(table, replacements)

def replace_everywhere(doc: Document, replacements: dict):
    replace_in_document_parts(doc, replacements)
    for section in doc.sections:
        replace_in_document_parts(section.header, replacements)
        replace_in_document_parts(section.footer, replacements)

def replace_map_placeholder_with_text(doc: Document, text: str, placeholder="{{Map}}"):
    def try_replace(paragraph):
        if placeholder in paragraph.text:
            full_text = "".join(run.text for run in paragraph.runs)
            new_text = full_text.replace(placeholder, text)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = new_text
            return True
        return False

    for paragraph in doc.paragraphs:
        if try_replace(paragraph):
            print("MAP PLACEHOLDER REPLACED WITH TEXT IN BODY")
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if try_replace(paragraph):
                        print("MAP PLACEHOLDER REPLACED WITH TEXT IN TABLE")
                        return True
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                if try_replace(paragraph):
                    print("MAP PLACEHOLDER REPLACED WITH TEXT IN HEADER/FOOTER")
                    return True
    print("MAP PLACEHOLDER NOT FOUND FOR TEXT REPLACEMENT")
    return False

def replace_map_placeholder(doc: Document, map_bytes: bytes, placeholder="{{Map}}"):
    def try_insert(paragraph):
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(map_bytes), width=Inches(MAP_WIDTH_INCHES), height=Inches(MAP_HEIGHT_INCHES))
            return True
        return False

    for paragraph in doc.paragraphs:
        if try_insert(paragraph):
            print("MAP INSERTED IN BODY")
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if try_insert(paragraph):
                        print("MAP INSERTED IN TABLE")
                        return True
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                if try_insert(paragraph):
                    print("MAP INSERTED IN HEADER/FOOTER")
                    return True
    print("MAP PLACEHOLDER NOT FOUND")
    return False

def build_replacements(data: dict) -> dict:
    replacements = {}
    for placeholder, col_id in COLUMN_MAP.items():
        replacements[placeholder] = data.get(col_id, "") or ""
    first_date = extract_first_date((data.get("dup__of_90__timeline", "") or "").strip()) or format_yyyy_mm_dd_to_dd_mm_yyyy((data.get("date_mm1rnmvg", "") or "").strip())
    replacements["{{Date}}"] = first_date
    replacements["{{Hit date}}"] = format_yyyy_mm_dd_to_dd_mm_yyyy((data.get("date_mm1rnmvg", "") or "").strip())
    replacements["{{date_today}}"] = datetime.now().strftime("%d.%m.%Y")
    replacements["{{ProjectName}}"] = data.get("name", "") or ""
    return replacements

def create_report(data: dict):
    report_type_value = data.get(REPORT_TYPE_COLUMN_ID, "")
    template_path = build_template_path(report_type_value)
    print("DOWNLOADING TEMPLATE FROM:", template_path)
    _, res = dbx.files_download(template_path)
    doc = Document(io.BytesIO(res.content))
    replacements = build_replacements(data)
    print("REPLACEMENTS:", replacements)
    replace_everywhere(doc, replacements)

    street = (data.get("text_mm12vcy9", "") or "").strip()
    number = (data.get("text_mm12jf0w", "") or "").strip()
    city = (data.get("text_mm264acy", "") or "").strip()

    map_bytes = None  # נשמר לשימוש חוזר בגרמושקא — ללא קריאה כפולה ל-Mapbox
    if street and number and city:
        try:
            map_bytes = create_mapbox_hybrid_image(street, number, city)
            replace_map_placeholder(doc, map_bytes, placeholder="{{Map}}")
        except Exception as e:
            print("MAP GENERATION FAILED:", str(e))
            replace_map_placeholder_with_text(doc, MAP_NOT_FOUND_TEXT, placeholder="{{Map}}")
    else:
        replace_map_placeholder_with_text(doc, MAP_NOT_FOUND_TEXT, placeholder="{{Map}}")

    buffer = io.BytesIO()
    doc.save(buffer)
    report_bytes = buffer.getvalue()
    print("REPORT CREATED IN MEMORY")
    return report_bytes, replacements, map_bytes  # map_bytes מועבר הלאה לגרמושקא

def process_item(item_id: int):
    print("START process_item with item_id:", item_id)
    data = get_item_data(item_id)
    print("ITEM DATA:", data)

    project_name = (data.get("name", "") or "").strip() or f"project_{item_id}"
    city_name = (data.get(CITY_COLUMN_ID, "") or "").strip()
    if not city_name:
        raise Exception(f"City field is empty or missing (column id: {CITY_COLUMN_ID})")

    report_folder = f"{BASE_REPORTS_PATH}/{city_name}/{project_name}"
    photos_folder = f"{report_folder}/תמונות"
    findings_folder = f"{report_folder}/ממצאים ראשוניים + כ.כמויות"

    street = (data.get("text_mm12vcy9", "") or "").strip()
    number = (data.get("text_mm12jf0w", "") or "").strip()
    apartment = (data.get("text_mm127a33", "") or "").strip()
    report_date = extract_first_date((data.get("dup__of_90__timeline", "") or "").strip()) or format_yyyy_mm_dd_to_dd_mm_yyyy((data.get("date_mm1rnmvg", "") or "").strip())

    file_name = sanitize_filename(f"מימצאים מבניים והנחיות ראשוניות רחוב {street} {number} - דירה {apartment}- {city_name}, {report_date}.docx")
    file_path = f"{findings_folder}/{file_name}"

    print("DROPBOX_SHARED_LINK:", DROPBOX_SHARED_LINK)
    print("SHARED_ROOT_PATH:", SHARED_ROOT_PATH)
    print("BASE_REPORTS_PATH:", BASE_REPORTS_PATH)
    print("TEMPLATE_DIR:", TEMPLATE_DIR)
    print("REPORT FOLDER:", report_folder)
    print("PHOTOS FOLDER:", photos_folder)
    print("FINDINGS FOLDER:", findings_folder)
    print("FILE PATH:", file_path)

    create_folder_if_needed(report_folder)
    create_folder_if_needed(photos_folder)
    create_folder_if_needed(findings_folder)

    report_bytes, replacements, map_bytes = create_report(data)

    dbx.files_upload(report_bytes, file_path, mode=dropbox.files.WriteMode.overwrite, mute=True)
    print("FILE UPLOADED TO DROPBOX")

    link = get_or_create_shared_link(file_path)

    try:
        upload_file_to_monday(item_id=item_id, column_id=FILES_COLUMN_ID, file_name=file_name, file_bytes=report_bytes)
        print("FILE UPLOADED TO MONDAY FILES COLUMN")
    except Exception as e:
        print("MONDAY FILE UPLOAD FAILED - CONTINUING:", str(e))

    update_link_column(item_id=item_id, column_id=LINK_COLUMN_ID, url=link, text=file_name)
    print("DROPBOX LINK WRITTEN TO MONDAY LINK COLUMN")

    # ── גרמושקא PDF ───────────────────────────────────────────────────────────
    # נוצרת מיד אחרי ה-Word, עם אותם נתונים ואותה תמונת מפה (ללא קריאה נוספת ל-Mapbox)
    try:
        garmushka_bytes = build_garmushka_pdf(data, map_bytes=map_bytes)
        garmushka_filename = sanitize_filename(
            f"גרמושקא רחוב {street} {number}"
            + (f" דירה {apartment}" if apartment else "")
            + f" {city_name}, {report_date}.pdf"
        )
        garmushka_path = f"{findings_folder}/{garmushka_filename}"

        dbx.files_upload(garmushka_bytes, garmushka_path, mode=dropbox.files.WriteMode.overwrite, mute=True)
        print("GARMUSHKA UPLOADED TO DROPBOX:", garmushka_path)

        try:
            upload_file_to_monday(
                item_id=item_id,
                column_id=FILES_COLUMN_ID,
                file_name=garmushka_filename,
                file_bytes=garmushka_bytes,
            )
            print("GARMUSHKA UPLOADED TO MONDAY FILES COLUMN")
        except Exception as e:
            print("GARMUSHKA MONDAY FILE UPLOAD FAILED - CONTINUING:", str(e))

    except Exception as e:
        # הגרמושקא לא מפילה את כל התהליך
        print("GARMUSHKA GENERATION FAILED - CONTINUING:", str(e))
    # ── /גרמושקא ─────────────────────────────────────────────────────────────

    print("PROCESS SUCCESS. LINK:", link)
    return {"link": link, "file_name": file_name, "replacements": replacements}

@app.route("/", methods=["GET"])
def home():
    return "OK", 200

@app.route("/webhook", methods=["GET", "POST"])
def webhook():
    if request.method == "GET":
        return "Webhook endpoint is live", 200
    data = request.get_json(silent=True) or {}
    print("INCOMING DATA:", data)
    if "challenge" in data:
        return jsonify({"challenge": data["challenge"]}), 200
    try:
        item_id = (data.get("event", {}).get("pulseId") or data.get("event", {}).get("itemId") or data.get("pulseId") or data.get("itemId"))
        print("ITEM ID:", item_id)
        if not item_id:
            return jsonify({"error": "No item id found", "payload": data}), 400
        result = process_item(int(item_id))
        return jsonify({"status": "success", **result}), 200
    except Exception as e:
        print("PROCESS ERROR:", str(e))
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
